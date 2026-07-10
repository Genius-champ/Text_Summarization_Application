"""
==========================================================
DOWNLOAD MODULE

Creates downloadable summaries in:

    - TXT format
    - DOCX format
    - PDF format

Libraries:
    - python-docx
    - reportlab
    - Streamlit

Compatible with:
    - Streamlit Cloud
    - Python 3.11+
==========================================================
"""

import io

import streamlit as st

from docx import Document

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import getSampleStyleSheet



###########################################################
# TXT DOWNLOAD
###########################################################

def download_txt(summary):
    """
    Creates a TXT download button.

    Parameters:
        summary (str)
    """

    txt_file = summary.encode(
        "utf-8"
    )


    st.download_button(

        label="📄 Download TXT",

        data=txt_file,

        file_name="summary.txt",

        mime="text/plain"

    )



###########################################################
# DOCX DOWNLOAD
###########################################################

def create_docx(summary):
    """
    Creates a Word document in memory.

    Returns:
        BytesIO object
    """

    document = Document()


    document.add_heading(
        "AI Generated Summary",
        level=1
    )


    document.add_paragraph(
        summary
    )


    buffer = io.BytesIO()


    document.save(
        buffer
    )


    buffer.seek(0)


    return buffer



def download_docx(summary):
    """
    Creates DOCX download button.
    """

    docx_file = create_docx(
        summary
    )


    st.download_button(

        label="📝 Download DOCX",

        data=docx_file,

        file_name="summary.docx",

        mime=
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    )



###########################################################
# PDF DOWNLOAD
###########################################################

def create_pdf(summary):
    """
    Creates a PDF file in memory.

    Returns:
        BytesIO object
    """

    buffer = io.BytesIO()


    pdf = SimpleDocTemplate(
        buffer
    )


    styles = getSampleStyleSheet()


    content = []


    title = Paragraph(
        "AI Generated Summary",
        styles["Title"]
    )


    content.append(
        title
    )


    content.append(
        Spacer(
            1,
            20
        )
    )


    paragraphs = summary.split(
        "\n"
    )


    for paragraph in paragraphs:

        if paragraph.strip():

            content.append(

                Paragraph(

                    paragraph,

                    styles["BodyText"]

                )

            )


            content.append(

                Spacer(
                    1,
                    12
                )

            )


    pdf.build(
        content
    )


    buffer.seek(0)


    return buffer



def download_pdf(summary):
    """
    Creates PDF download button.
    """

    pdf_file = create_pdf(
        summary
    )


    st.download_button(

        label="📕 Download PDF",

        data=pdf_file,

        file_name="summary.pdf",

        mime="application/pdf"

    )
